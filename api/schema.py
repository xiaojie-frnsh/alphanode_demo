import graphene
from graphene_django import DjangoObjectType
from graphene_file_upload.scalars import Upload
from .models import Document, DocumentQuery
import os
from django.conf import settings
from openai import OpenAI
import PyPDF2
import docx
import io
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import logging
from tenacity import retry, stop_after_attempt, wait_exponential
import time
from typing import List
import json
from datetime import datetime
from pathlib import Path
from .callbacks import LLMLoggingCallbackHandler
from typing import Optional
from asgiref.sync import sync_to_async
from functools import partial
from django.db import transaction

logger = logging.getLogger(__name__)

def log_llm_interaction(question, prompt, response, document_id):
    """Log LLM interactions to a file"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_dir = Path("logs/llm_interactions")
    log_dir.mkdir(parents=True, exist_ok=True)
    
    interaction_data = {
        "timestamp": timestamp,
        "document_id": document_id,
        "question": question,
        "prompt": prompt,
        "response": response
    }
    
    log_file = log_dir / f"interaction_{timestamp}_{document_id}.json"
    with open(log_file, "w", encoding="utf-8") as f:
        json.dump(interaction_data, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Logged LLM interaction to {log_file}")

class DocumentType(DjangoObjectType):
    class Meta:
        model = Document
        fields = '__all__'

class DocumentQueryType(DjangoObjectType):
    class Meta:
        model = DocumentQuery
        fields = ('id', 'question', 'answer', 'created_at', 'document')

# Add retry decorator for OpenAI API calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    reraise=True
)
def get_embeddings(embeddings_client, text):
    """Wrapper function to handle retries for embeddings"""
    logger.debug(f"Attempting to get embeddings for text of length: {len(text)}")
    return embeddings_client.embed_documents([text])[0]

def process_document_content(document):
    """Extract text content and create vector embeddings."""
    logger.info(f"Starting document processing for document ID: {document.id}")
    file_path = document.file.path
    content = ""
    
    try:
        if document.file_type.lower() == 'pdf':
            logger.debug(f"Processing PDF file: {file_path}")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                for i, page in enumerate(pdf_reader.pages):
                    logger.debug(f"Extracting text from page {i+1}")
                    content += page.extract_text() + "\n"
        
        elif document.file_type.lower() in ['docx', 'doc']:
            logger.debug(f"Processing Word document: {file_path}")
            doc = docx.Document(file_path)
            logger.info(f"Word document has {len(doc.paragraphs)} paragraphs")
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        elif document.file_type.lower() == 'txt':
            logger.debug(f"Processing text file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        
        logger.info(f"Successfully extracted {len(content)} characters of text")
        
        # Split the content into chunks
        logger.debug("Starting text splitting process")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # Increased to capture more context
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # More granular separators
        )
        chunks = text_splitter.split_text(content)
        logger.info(f"Split text into {len(chunks)} chunks")

        # Create a unique collection name and ensure the directory exists
        collection_name = f"doc_{document.id}"
        persist_directory = os.path.join(settings.VECTOR_STORE_DIR, collection_name)
        
        # Ensure the vector store directory exists
        os.makedirs(settings.VECTOR_STORE_DIR, exist_ok=True)
        logger.info(f"Creating vector store at: {persist_directory}")
        
        try:
            # Create embeddings
            embeddings = OpenAIEmbeddings(
                openai_api_key=settings.OPENAI_API_KEY,
                max_retries=3,
                timeout=30
            )
            
            # Store the vectors
            vectordb = Chroma.from_texts(
                texts=chunks,
                embedding=embeddings,
                persist_directory=persist_directory
            )
            vectordb.persist()
            logger.info(f"Successfully stored vectors in ChromaDB at {persist_directory}")
            
            return content, persist_directory
            
        except Exception as e:
            logger.error(f"Error creating vector store: {str(e)}", exc_info=True)
            raise

    except Exception as e:
        logger.error(f"Error processing document: {str(e)}", exc_info=True)
        raise

class UploadDocumentMutation(graphene.Mutation):
    class Arguments:
        file = Upload(required=True)
        title = graphene.String(required=True)

    success = graphene.Boolean()
    document = graphene.Field(DocumentType)

    def mutate(self, info, file, title):
        try:
            document = Document.objects.create(
                title=title,
                file=file,
                file_type=file.name.split('.')[-1],
                content=''  # Empty content for now
            )
            return UploadDocumentMutation(success=True, document=document)
        except Exception as e:
            return UploadDocumentMutation(success=False, document=None)

class AskDocumentResponse(graphene.ObjectType):
    success = graphene.Boolean()
    query = graphene.Field(DocumentQueryType)
    error = graphene.String()

class AskDocumentMutation(graphene.Mutation):
    class Arguments:
        document_id = graphene.ID(required=True)
        question = graphene.String(required=True)

    Output = AskDocumentResponse

    async def mutate(self, info, document_id, question):
        logger.info(f"Processing question for document ID: {document_id}")
        logger.debug(f"Question: {question}")
        
        try:
            # Make database operations async-compatible
            get_document = sync_to_async(Document.objects.get)
            document = await get_document(id=document_id)
            
            # Check if vector_store_path is None or if the path doesn't exist
            if not document.vector_store_path or not os.path.exists(document.vector_store_path):
                logger.info("Vector store path is missing or invalid, processing document...")
                try:
                    content, vector_store_path = process_document_content(document)
                    document.content = content
                    document.vector_store_path = vector_store_path
                    document.save()
                    logger.info(f"Document processed and saved with vector store path: {vector_store_path}")
                except Exception as e:
                    logger.error(f"Error processing document content: {str(e)}", exc_info=True)
                    raise
            
            # Verify vector store path after processing
            if not document.vector_store_path:
                logger.error("Vector store path is still None after processing")
                raise ValueError("Failed to create vector store path")
            
            # Initialize RAG components
            try:
                embeddings = OpenAIEmbeddings(
                    openai_api_key=settings.OPENAI_API_KEY,
                    max_retries=3,
                    timeout=30
                )
                vectordb = Chroma(
                    persist_directory=document.vector_store_path,
                    embedding_function=embeddings
                )
                logger.info("Successfully initialized vector store")
            except Exception as e:
                logger.error(f"Error initializing vector store: {str(e)}", exc_info=True)
                raise

            # Update retriever configuration - removing unsupported parameters
            retriever = vectordb.as_retriever(
                search_type="similarity",
                search_kwargs={
                    "k": 16,  # Number of documents to retrieve
                }
            )

            # Enhanced prompt template with stronger emphasis on document title and introduction
            prompt_template = """You are a helpful business analyst analyzing documents and reports. Use the following pieces of context to answer the question.

            Context: {context}

            Question: {question}

            Instructions:
            1. Focus on the main content at the beginning of the document, not the footer or administrative sections
            2. For topic questions, include:
               - Document title
               - Main subject matter
               - Key statistics from the first few paragraphs
            3. Ignore boilerplate text about websites, subscriptions, or administrative information
            4. Be specific and concise
            5. Only use information directly stated in the context
            6. Make sure you output human readable answer

            Answer:"""
            
            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )

            # Update LLM configuration to use GPT-4
            llm = ChatOpenAI(
                openai_api_key=settings.OPENAI_API_KEY,
                model_name="gpt-4",  # Changed from gpt-3.5-turbo to gpt-4
                temperature=0,  # Keeping temperature at 0 for consistent, factual responses
                max_retries=3,
                timeout=30
            )

            # Create the RAG chain with the updated components
            callback_handler = LLMLoggingCallbackHandler()

            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=retriever,
                return_source_documents=True,
                callbacks=[callback_handler],
                chain_type_kwargs={
                    "prompt": PROMPT,
                    "verbose": True
                }
            )

            # Get the answer with logging
            logger.info("Generating answer")
            try:
                # Make retriever operations async-compatible
                get_relevant_docs = sync_to_async(retriever.get_relevant_documents)
                retrieved_docs = await get_relevant_docs(question)
                retrieved_context = "\n\n".join(doc.page_content for doc in retrieved_docs)
                
                # Format prompt and get GPT-4 answer
                full_prompt = PROMPT.format(
                    context=retrieved_context,
                    question=question
                )
                
                qa_chain_call = sync_to_async(qa_chain)
                result = await qa_chain_call({"query": question})
                answer = result["result"]
                
                # Make logging async-compatible
                log_interaction = sync_to_async(log_llm_interaction)
                await log_interaction(
                    question=question,
                    prompt=full_prompt,
                    response={
                        "gpt4_answer": answer,
                        "perplexity_answer": None
                    },
                    document_id=document_id
                )
                
                # Log retrieved documents for debugging
                if "source_documents" in result:
                    logger.debug(f"Retrieved {len(result['source_documents'])} chunks:")
                    for i, doc in enumerate(result["source_documents"]):
                        logger.debug(f"Chunk {i+1} (first 200 chars): {doc.page_content[:200]}")
                        logger.debug(f"Chunk {i+1} metadata: {doc.metadata}")
            
            except Exception as e:
                logger.error(f"Error generating answer: {str(e)}", exc_info=True)
                raise

            # Wrap the create operation in sync_to_async
            @sync_to_async
            @transaction.atomic
            def create_query():
                return DocumentQuery.objects.create(
                    question=question,
                    answer=answer,
                    document=document
                )
            
            # Create the query asynchronously
            query = await create_query()

            return AskDocumentResponse(
                success=True,
                query=query,
                error=None
            )
            
        except Exception as e:
            logger.error(f"Error in askDocument mutation: {str(e)}", exc_info=True)
            return AskDocumentResponse(
                success=False,
                query=None,
                error=str(e)
            )

class Mutation(graphene.ObjectType):
    upload_document = UploadDocumentMutation.Field()
    ask_document = AskDocumentMutation.Field()

class Query(graphene.ObjectType):
    all_documents = graphene.List(DocumentType)
    all_queries = graphene.List(DocumentQueryType)
    document_queries = graphene.List(DocumentQueryType, document_id=graphene.ID(required=True))

    def resolve_all_documents(self, info):
        return Document.objects.all()

    def resolve_all_queries(self, info):
        return DocumentQuery.objects.all()

    def resolve_document_queries(self, info, document_id):
        return DocumentQuery.objects.filter(document_id=document_id)

schema = graphene.Schema(query=Query, mutation=Mutation) 